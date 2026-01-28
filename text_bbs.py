import time
import os
import json
from playwright.sync_api import sync_playwright
from docx import Document

# --- 核心配置 ---
PROGRESS_FILE = "crawl_progress.json"  # 记录已抓取的链接，防止置顶帖重复
SAVE_INTERVAL = 100                    # 每 100 篇帖子保存一个 Word 文件
DOC_DIR = "WuShu_Archive"              # 结果存放文件夹

if not os.path.exists(DOC_DIR):
    os.makedirs(DOC_DIR)

def load_progress():
    if os.path.exists(PROGRESS_FILE):
        with open(PROGRESS_FILE, 'r') as f:
            data = json.load(f)
            return set(data.get("seen_hrefs", [])), data.get("total_count", 0), data.get("file_idx", 1)
    return set(), 0, 1

def save_progress(seen_hrefs, total_count, file_idx):
    with open(PROGRESS_FILE, 'w') as f:
        json.dump({
            "seen_hrefs": list(seen_hrefs),
            "total_count": total_count,
            "file_idx": file_idx
        }, f)

def run():
    seen_hrefs, total_count, file_idx = load_progress()

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=False) # 必须可见，方便手动操作
        context = browser.new_context()
        page = context.new_page()

        print(f">>> 准备开始！历史已处理 {total_count} 篇。")
        page.goto("https://wvpn.ustc.edu.cn/")
        
        print("\n" + "="*50)
        print("【手动操作步骤】")
        print("1. 登录 WebVPN。")
        print("2. 进入精武门 (WuShu) 板块。")
        print("3. 【关键】：点击底部的 [第一页] 按钮。")
        print("4. 确认列表序号从 1 开始后，回到这里按【回车】。")
        print("="*50 + "\n")
        
        input(">>> 确认已在第一页？按回车启动自动抓取...")

        f3 = page.frame_locator('frame[name="f3"]')
        doc = Document()

        try:
            while True:
                # 1. 扫描当前页帖子
                f3.locator("a.o_title").first.wait_for(state="visible", timeout=20000)
                links = f3.locator("a.o_title").all()
                
                posts_on_page = []
                for l in links:
                    href = l.get_attribute("href")
                    title = l.inner_text().strip()
                    if href:
                        posts_on_page.append((title, href, l))

                print(f"\n>>> 正在扫描本页 ({len(posts_on_page)} 条链接)...")

                # 2. 逐一抓取（带去重）
                for title, href, handle in posts_on_page:
                    if href in seen_hrefs:
                        continue # 跳过置顶帖或已抓取帖

                    try:
                        # Ctrl+点击打开新标签页，绕过 WebVPN 路径转换 Bug
                        with context.expect_page() as new_page_info:
                            handle.click(modifiers=["Control"])
                        
                        post_page = new_page_info.value
                        post_page.wait_for_load_state("domcontentloaded")
                        
                        # 提取正文 (根据你的截图 image_ca7a47)
                        content_el = post_page.locator("div.post_text")
                        content_el.wait_for(timeout=5000)
                        text = content_el.inner_text()
                        
                        doc.add_heading(title, level=1)
                        doc.add_paragraph(text)
                        doc.add_page_break()
                        
                        seen_hrefs.add(href)
                        total_count += 1
                        print(f"  [{total_count}] 成功: {title[:15]}...")
                        
                        post_page.close()
                        time.sleep(0.4) # 稍微休息下，毕竟 7000 多篇呢

                        # 达到保存间隔
                        if total_count % SAVE_INTERVAL == 0:
                            save_path = os.path.join(DOC_DIR, f"WuShu_Part_{file_idx}.docx")
                            doc.save(save_path)
                            file_idx += 1
                            save_progress(seen_hrefs, total_count, file_idx)
                            print(f"--- 进度存档: {save_path} ---")
                            doc = Document() # 开启新文档

                    except Exception as e:
                        print(f"  [!] 跳过帖子 {title}: {e}")
                        if 'post_page' in locals(): post_page.close()

                # 3. 翻页逻辑 (基于 image_c936de)
                next_btn = f3.locator("a.next").first
                # 检查是否有 href 属性且 class 不含 disabled
                btn_class = next_btn.get_attribute("class") or ""
                has_href = next_btn.get_attribute("href") is not None
                
                if next_btn.is_visible() and has_href and "disabled" not in btn_class:
                    print(">>> 正在翻向下一页 (Next)...")
                    next_btn.click()
                    time.sleep(3) # 给框架一点加载时间
                    save_progress(seen_hrefs, total_count, file_idx) # 翻页后存一下进度
                else:
                    print(">>> 探测到最后一页或翻页按钮失效。")
                    break

        except Exception as e:
            print(f"\n>>> 发生意外中断: {e}")
        finally:
            # 无论如何，保存最后的劳动成果
            final_path = os.path.join(DOC_DIR, f"WuShu_Final_Part.docx")
            doc.save(final_path)
            save_progress(seen_hrefs, total_count, file_idx)
            print(f">>> 最终进度已保存，共处理 {total_count} 篇。")
            browser.close()

if __name__ == "__main__":
    run()